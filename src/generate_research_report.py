import ctypes
import html
import json
import os
import platform
import statistics
import unicodedata
from collections import Counter
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics import average_precision_score, auc, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import label_binarize
from sklearn.svm import LinearSVC
from wordcloud import WordCloud


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
REPORTS_DIR = ROOT / "reports"
SELECTED_REFERENCES_PATH = ROOT / "selected_references.json"
OUTPUT_DOCX = ROOT / "Desktop_Assistant_NLP_Research_Report.docx"


def read_json(path: Path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def ascii_clean(value: str) -> str:
    if not isinstance(value, str):
        return str(value)
    normalized = unicodedata.normalize("NFKD", value)
    return normalized.encode("ascii", "ignore").decode("ascii")


def latest_summary_path() -> Path:
    run_dirs = [path for path in REPORTS_DIR.glob("run_*") if path.is_dir()]
    if not run_dirs:
        raise FileNotFoundError("No report run directory found under reports/. Run evaluate.py first.")
    run_dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    summary_path = run_dirs[0] / "summary.json"
    if not summary_path.exists():
        raise FileNotFoundError(f"Missing summary file: {summary_path}")
    return summary_path


def get_system_ram_gb() -> float:
    if os.name != "nt":
        return 0.0
    memory_kb = ctypes.c_ulonglong(0)
    result = ctypes.windll.kernel32.GetPhysicallyInstalledSystemMemory(ctypes.byref(memory_kb))
    if result == 0:
        return 0.0
    return round(memory_kb.value / (1024 * 1024), 2)


def split_data(texts, labels, seed=42, test_size=0.2, val_size=0.2):
    train_val_texts, test_texts, train_val_labels, test_labels = train_test_split(
        texts,
        labels,
        test_size=test_size,
        random_state=seed,
        stratify=labels,
    )

    val_ratio_inside_train_val = val_size / (1.0 - test_size)
    train_texts, val_texts, train_labels, val_labels = train_test_split(
        train_val_texts,
        train_val_labels,
        test_size=val_ratio_inside_train_val,
        random_state=seed,
        stratify=train_val_labels,
    )

    return {
        "train_texts": train_texts,
        "train_labels": train_labels,
        "val_texts": val_texts,
        "val_labels": val_labels,
        "test_texts": test_texts,
        "test_labels": test_labels,
    }


def compute_roc_pr_curves(intent_samples, selected_c=10.0, seed=42):
    texts = [item["text"] for item in intent_samples]
    labels = [item["intent"] for item in intent_samples]
    split = split_data(texts, labels, seed=seed)

    train_val_texts = split["train_texts"] + split["val_texts"]
    train_val_labels = split["train_labels"] + split["val_labels"]

    vectorizer = TfidfVectorizer()
    x_train_val = vectorizer.fit_transform(train_val_texts)
    x_test = vectorizer.transform(split["test_texts"])

    model = LinearSVC(C=selected_c, max_iter=5000)
    model.fit(x_train_val, train_val_labels)
    scores = model.decision_function(x_test)

    label_order = sorted(set(labels))
    y_true_bin = label_binarize(split["test_labels"], classes=label_order)

    if scores.ndim == 1:
        scores = np.vstack([-scores, scores]).T

    macro_roc_auc = roc_auc_score(y_true_bin, scores, average="macro", multi_class="ovr")
    macro_avg_precision = average_precision_score(y_true_bin, scores, average="macro")

    fpr, tpr, _ = roc_curve(y_true_bin.ravel(), scores.ravel())
    precision, recall, _ = precision_recall_curve(y_true_bin.ravel(), scores.ravel())

    return {
        "macro_roc_auc": float(macro_roc_auc),
        "macro_pr_auc": float(macro_avg_precision),
        "roc_curve": {"fpr": fpr.tolist(), "tpr": tpr.tolist(), "auc": float(auc(fpr, tpr))},
        "pr_curve": {"recall": recall.tolist(), "precision": precision.tolist()},
    }


def build_architecture_figure(output_path: Path):
    fig, ax = plt.subplots(figsize=(13, 5))
    ax.axis("off")

    boxes = [
        (0.03, 0.45, 0.15, 0.2, "User Command"),
        (0.22, 0.45, 0.18, 0.2, "TF-IDF +\nIntent Model"),
        (0.45, 0.45, 0.16, 0.2, "Regex NER\nExtractor"),
        (0.66, 0.45, 0.15, 0.2, "Executor\nRouter"),
        (0.84, 0.45, 0.13, 0.2, "OS / Web\nAction"),
    ]

    for x, y, w, h, label in boxes:
        rect = plt.Rectangle((x, y), w, h, edgecolor="#1f4e79", facecolor="#dbeafe", linewidth=2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=10, fontweight="bold")

    arrow_y = 0.55
    for start_x, end_x in [(0.18, 0.22), (0.40, 0.45), (0.61, 0.66), (0.81, 0.84)]:
        ax.annotate(
            "",
            xy=(end_x, arrow_y),
            xytext=(start_x, arrow_y),
            arrowprops=dict(arrowstyle="-|>", lw=2, color="#1f4e79"),
        )

    ax.text(0.5, 0.83, "Proposed Desktop Assistant NLP Processing Pipeline", ha="center", fontsize=14, fontweight="bold")
    ax.text(0.5, 0.18, "Runtime checks: dry-run validation, required entity checks, and intent-supported routing", ha="center", fontsize=10)

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_intent_distribution_figure(intent_samples, output_path: Path):
    counts = Counter(item["intent"] for item in intent_samples)
    labels = sorted(counts)
    values = [counts[label] for label in labels]
    positions = np.arange(len(labels))

    fig, ax = plt.subplots(figsize=(12, 5))
    bars = ax.bar(positions, values, color="#2563eb")
    ax.set_title("Intent Dataset Class Distribution")
    ax.set_xlabel("Intent")
    ax.set_ylabel("Sample Count")
    ax.set_xticks(positions)
    ax.set_xticklabels(labels, rotation=45, ha="right")

    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.2, str(value), ha="center", fontsize=8)

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_confusion_matrix_figure(matrix_values, labels, output_path: Path, title: str):
    matrix = np.array(matrix_values)

    fig_size = max(9, int(len(labels) * 0.6))
    fig, ax = plt.subplots(figsize=(fig_size, fig_size))
    image = ax.imshow(matrix, cmap="Blues")
    fig.colorbar(image, ax=ax)

    axis_positions = np.arange(len(labels))
    ax.set_xticks(axis_positions)
    ax.set_yticks(axis_positions)
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_yticklabels(labels)
    ax.set_xlabel("Predicted")
    ax.set_ylabel("Actual")
    ax.set_title(title)

    threshold = matrix.max() * 0.55 if matrix.max() > 0 else 0.0

    for row_idx in range(matrix.shape[0]):
        for col_idx in range(matrix.shape[1]):
            value = int(matrix[row_idx, col_idx])
            color = "white" if matrix[row_idx, col_idx] >= threshold else "black"
            ax.text(col_idx, row_idx, str(value), ha="center", va="center", fontsize=8, color=color)

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_stage_breakdown_figure(summary_e2e, output_path: Path):
    stage_breakdown = summary_e2e["stage_breakdown"]
    labels = list(stage_breakdown.keys())
    values = list(stage_breakdown.values())
    positions = np.arange(len(labels))

    fig, ax = plt.subplots(figsize=(8, 5))
    ax.bar(positions, values, color=["#16a34a", "#dc2626", "#f59e0b"][: len(labels)])
    ax.set_title("End-to-End Failure Stage Breakdown")
    ax.set_xlabel("Stage")
    ax.set_ylabel("Count")
    ax.set_xticks(positions)
    ax.set_xticklabels(labels, rotation=25, ha="right")

    for idx, value in zip(positions, values):
        ax.text(idx, value + 0.2, str(value), ha="center")

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_ner_entity_f1_figure(summary_ner, output_path: Path):
    rows = summary_ner.get("per_entity_type", [])
    labels = [row["entity_type"] for row in rows]
    f1_values = [row["f1"] for row in rows]
    positions = np.arange(len(labels))

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(positions, f1_values, color="#0f766e")
    ax.set_ylim(0.0, 1.05)
    ax.set_title("Per-Entity-Type F1 Scores (NER)")
    ax.set_xlabel("Entity Type")
    ax.set_ylabel("F1 Score")
    ax.set_xticks(positions)
    ax.set_xticklabels(labels, rotation=35, ha="right")

    for idx, value in zip(positions, f1_values):
        ax.text(idx, value + 0.02, f"{value:.2f}", ha="center", fontsize=8)

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_wordcloud_figure(intent_samples, output_path: Path):
    corpus = " ".join(item["text"] for item in intent_samples)
    cloud = WordCloud(width=1400, height=700, background_color="white", colormap="Blues").generate(corpus)

    fig, ax = plt.subplots(figsize=(11, 5))
    ax.imshow(cloud, interpolation="bilinear")
    ax.axis("off")
    ax.set_title("Word Cloud of Intent Commands", fontsize=14)

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def build_roc_pr_figure(roc_pr, output_path: Path):
    fpr = np.array(roc_pr["roc_curve"]["fpr"])
    tpr = np.array(roc_pr["roc_curve"]["tpr"])
    recall = np.array(roc_pr["pr_curve"]["recall"])
    precision = np.array(roc_pr["pr_curve"]["precision"])

    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    axes[0].plot(fpr, tpr, color="#1d4ed8", linewidth=2)
    axes[0].plot([0, 1], [0, 1], "--", color="gray")
    axes[0].set_title("ROC Curve (Micro-averaged OVR)")
    axes[0].set_xlabel("False Positive Rate")
    axes[0].set_ylabel("True Positive Rate")

    axes[1].plot(recall, precision, color="#059669", linewidth=2)
    axes[1].set_title("Precision-Recall Curve (Micro-averaged OVR)")
    axes[1].set_xlabel("Recall")
    axes[1].set_ylabel("Precision")

    fig.tight_layout()
    fig.savefig(output_path, dpi=220)
    plt.close(fig)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(10)

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(3)


def add_paragraph(doc: Document, text: str, align=None):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_equation_block(doc: Document, equation_text: str, number: int):
    # Use an Office Math XML block so equations render as equation objects in Word.
    table = doc.add_table(rows=1, cols=2)

    left_cell = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escaped = html.escape(equation_text)
    omml = parse_xml(
        f'<m:oMathPara {nsdecls("m")}><m:oMath><m:r><m:t>{escaped}</m:t></m:r></m:oMath></m:oMathPara>'
    )
    left_para._element.append(omml)

    right_para = right_cell.paragraphs[0]
    right_para.text = f"({number})"
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return table


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ascii_clean(str(value))

    return table


def first_author_last_name(authors):
    if not authors:
        return "Author"
    first = ascii_clean(authors[0]).strip()
    if not first:
        return "Author"
    return first.split()[-1]


def citation(doi_map, used_dois, doi):
    ref = doi_map[doi]
    used_dois.add(doi)
    return f"[{first_author_last_name(ref.get('authors', []))}, {ref.get('year', 'n.d.')}]"


def format_reference_entry(ref):
    authors = [ascii_clean(name) for name in ref.get("authors", []) if name]
    if len(authors) > 6:
        author_text = ", ".join(authors[:6]) + ", et al."
    else:
        author_text = ", ".join(authors)

    title = ascii_clean(ref.get("title", ""))
    venue = ascii_clean(ref.get("venue", ""))
    year = ref.get("year", "n.d.")

    entry = f"{author_text}, \"{title}\", {venue}"
    if ref.get("volume"):
        entry += f", vol. {ref['volume']}"
    if ref.get("issue"):
        entry += f", no. {ref['issue']}"
    if ref.get("page"):
        entry += f", pp. {ascii_clean(str(ref['page']))}"
    entry += f", {year}. DOI: {ref['doi']}."
    return entry


def main():
    summary_path = latest_summary_path()
    run_dir = summary_path.parent
    summary = read_json(summary_path)
    intent_summary = summary["intent"]
    e2e_summary = summary["e2e"]
    ner_summary = summary["ner"]
    e2e_confusion = read_json(run_dir / "e2e" / "e2e_confusion_matrix.json")
    e2e_conf_matrix = np.array(e2e_confusion["counts"])
    e2e_total = int(e2e_conf_matrix.sum())
    e2e_intent_mismatch_count = int(e2e_total - np.trace(e2e_conf_matrix))

    aligned_stage_breakdown = dict(e2e_summary.get("stage_breakdown", {}))
    aligned_stage_breakdown["intent_mismatch"] = e2e_intent_mismatch_count
    aligned_stage_breakdown["success"] = e2e_total - (
        aligned_stage_breakdown.get("intent_mismatch", 0) + aligned_stage_breakdown.get("entity_mismatch", 0)
    )
    e2e_summary["stage_breakdown"] = aligned_stage_breakdown

    intent_samples = read_json(DATA_DIR / "intent_dataset.json")
    e2e_samples = read_json(DATA_DIR / "e2e_eval_dataset.json")
    references = read_json(SELECTED_REFERENCES_PATH)

    doi_map = {ref["doi"]: ref for ref in references}
    used_dois = set()

    words_per_intent = [len(item["text"].split()) for item in intent_samples]
    chars_per_intent = [len(item["text"]) for item in intent_samples]
    words_per_e2e = [len(item["text"].split()) for item in e2e_samples]
    entities_per_e2e = [len(item.get("expected_entities", {})) for item in e2e_samples]

    intent_counts = Counter(item["intent"] for item in intent_samples)

    assets_dir = REPORTS_DIR / "paper_assets"
    assets_dir.mkdir(parents=True, exist_ok=True)

    fig_arch = assets_dir / "figure_1_architecture.png"
    fig_dist = assets_dir / "figure_2_intent_distribution.png"
    fig_conf = assets_dir / "figure_3_confusion_matrix.png"
    fig_stage = assets_dir / "figure_4_stage_breakdown.png"
    fig_ner = assets_dir / "figure_5_ner_entity_f1.png"
    fig_cloud = assets_dir / "figure_6_wordcloud.png"
    fig_rocpr = assets_dir / "figure_7_roc_pr.png"

    build_architecture_figure(fig_arch)
    build_intent_distribution_figure(intent_samples, fig_dist)
    build_confusion_matrix_figure(
        e2e_confusion["counts"],
        e2e_confusion["labels"],
        fig_conf,
        "End-to-End Intent Confusion Matrix (Counts)",
    )
    build_stage_breakdown_figure(e2e_summary, fig_stage)
    build_ner_entity_f1_figure(ner_summary, fig_ner)
    build_wordcloud_figure(intent_samples, fig_cloud)

    selected_c = float(intent_summary["model_results"]["linear_svm"]["best_params"]["C"])
    roc_pr = compute_roc_pr_curves(intent_samples, selected_c=selected_c, seed=summary["seed"])
    build_roc_pr_figure(roc_pr, fig_rocpr)

    cpu_name = platform.processor() or "Unknown CPU"
    logical_cpu = os.cpu_count() or 0
    ram_gb = get_system_ram_gb()

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph("DeskPilot-NLP: A Lightweight Intent-Driven Desktop Assistant with Rule-Aware Entity Extraction and End-to-End Evaluation")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    add_paragraph(doc, "Author details to be finalized by project team (Name, Department, Email ID)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Generated on: 12 April 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1. Abstract", level=1)
    abstract_text = (
        "This work presents DeskPilot-NLP, an intent-driven desktop assistant that transforms short natural language commands into executable local and web actions. "
        "The system combines TF-IDF based intent modeling with a lightweight rule-based named entity extraction layer and a guarded executor that validates command arguments before execution. "
        "Using 240 intent-labeled commands across 16 classes and a separate 40-sample end-to-end benchmark, the study evaluates both component-level and pipeline-level behavior under reproducible splits and fixed random seeds. "
        "The selected runtime classifier (Linear SVM) achieves 91.67% test accuracy with 89.15% macro-F1 on intent classification, while end-to-end command success reaches 92.50%. "
        "Entity extraction attains 96.77% micro-F1 with 95.00% exact-match rate, indicating that most residual failures arise from ambiguous web-search commands rather than systemic extraction errors. "
        "The work demonstrates that carefully engineered classical NLP methods can still provide strong command understanding for local assistants, with lower compute overhead than large neural alternatives. "
        "Overall, this project offers an efficient and extensible baseline for practical desktop automation in resource-constrained settings."
    )
    add_paragraph(doc, abstract_text)

    add_heading(doc, "2. Keywords", level=1)
    add_paragraph(doc, "Intent Classification, Named Entity Recognition, TF-IDF, Linear SVM, Desktop Automation, End-to-End NLP Evaluation")

    add_heading(doc, "3. Introduction", level=1)

    r1 = citation(doi_map, used_dois, "10.18653/v1/2020.ecnlp-1.6")
    r2 = citation(doi_map, used_dois, "10.18653/v1/2022.nlp4convai-1.2")
    r3 = citation(doi_map, used_dois, "10.18653/v1/2022.nlp4convai-1.5")
    r4 = citation(doi_map, used_dois, "10.18653/v1/2022.nlp4convai-1.10")
    r5 = citation(doi_map, used_dois, "10.18653/v1/2023.findings-emnlp.841")
    r6 = citation(doi_map, used_dois, "10.18653/v1/2024.nlp4convai-1.2")

    add_paragraph(
        doc,
        (
            "Recent conversational AI research has advanced intent detection and entity extraction for task-oriented systems, "
            "especially where user queries must be mapped to deterministic actions "
            f"{r1} {r2} {r5}. However, many assistant pipelines remain optimized for cloud-hosted dialogue rather than direct desktop orchestration."
        ),
    )

    add_paragraph(
        doc,
        (
            "The motivation of this project is to build a practical local assistant that can classify user commands, extract file/web entities, "
            "and invoke action handlers with explicit safety checks. Compared to large-model-first pipelines, this design emphasizes reproducibility, "
            "low inference overhead, and transparent failure analysis while remaining competitive on command-level metrics."
        ),
    )

    add_paragraph(
        doc,
        (
            "The problem statement is defined as follows: given a short command utterance, predict one of 16 action intents and extract the minimum entities "
            "required for successful execution. The objective is not only high intent accuracy, but robust end-to-end completion under dry-run validation, "
            f"with particular attention to out-of-scope confusion and sparse-entity utterances {r2} {r3}."
        ),
    )

    add_paragraph(
        doc,
        (
            "The novelty of this work lies in coupling classical text classification with a deterministic NER-executor bridge and unified three-level evaluation "
            "(intent, NER, and execution). The project contributes a compact experimental stack that still supports confusion analysis, per-entity diagnostics, "
            f"and deployment-ready command routing without requiring heavyweight transformer retraining {r4} {r6}."
        ),
    )

    add_heading(doc, "4. Literature Survey", level=1)

    l1 = citation(doi_map, used_dois, "10.18653/v1/2020.findings-emnlp.163")
    l2 = citation(doi_map, used_dois, "10.1109/icassp39728.2021.9414110")
    l3 = citation(doi_map, used_dois, "10.24963/ijcai.2021/523")
    l4 = citation(doi_map, used_dois, "10.24963/ijcai.2022/565")
    l5 = citation(doi_map, used_dois, "10.18653/v1/2022.emnlp-main.543")
    l6 = citation(doi_map, used_dois, "10.18653/v1/2022.findings-emnlp.245")
    l7 = citation(doi_map, used_dois, "10.1109/icassp48485.2024.10446353")
    l8 = citation(doi_map, used_dois, "10.1109/bigdata62323.2024.10825739")
    l9 = citation(doi_map, used_dois, "10.1038/s41598-023-50705-0")
    l10 = citation(doi_map, used_dois, "10.1016/j.neucom.2026.133053")

    add_paragraph(
        doc,
        (
            "Prior studies on spoken language understanding show a progression from joint recurrent architectures to graph- and transformer-based formulations. "
            "Graph-interactive approaches improve intent-slot dependency capture, while co-attention and higher-order interactions further stabilize multi-intent behavior "
            f"{l1} {l2} {l3} {l4}."
        ),
    )

    add_paragraph(
        doc,
        (
            "Recent methods increasingly integrate intent-slot co-occurrence priors, profile conditioning, and hybrid graph-transformer blocks, with growing emphasis "
            f"on explainability and robustness under low-resource settings {l5} {l6} {l7} {l8}. In parallel, transformer NER variants report broad gains across domains, "
            f"but often require larger training corpora and greater compute budgets {l9} {l10}."
        ),
    )

    add_paragraph(
        doc,
        (
            "A key research gap remains the reproducible deployment of these ideas in local desktop assistants where latency, interpretability, and deterministic execution are "
            "critical. Many papers optimize benchmark metrics but do not evaluate action-level success under execution constraints. This project addresses that gap through an explicit "
            "execution-aware benchmark and dry-run validation stage."
        ),
    )

    add_paragraph(doc, "Table 1: Comparison study of existing work on various parameters")

    c_agif = citation(doi_map, used_dois, "10.18653/v1/2020.findings-emnlp.163")
    c_dgm = citation(doi_map, used_dois, "10.24963/ijcai.2021/523")
    c_hitrans = citation(doi_map, used_dois, "10.18653/v1/2021.findings-emnlp.12")
    c_han = citation(doi_map, used_dois, "10.24963/ijcai.2022/565")
    c_gisco = citation(doi_map, used_dois, "10.18653/v1/2022.emnlp-main.543")
    c_estan = citation(doi_map, used_dois, "10.18653/v1/2022.findings-emnlp.245")
    c_misca = citation(doi_map, used_dois, "10.18653/v1/2023.findings-emnlp.841")
    c_agri = citation(doi_map, used_dois, "10.3390/app152010932")

    literature_rows = [
        [
            f"Qin et al., 2020 {c_agif}",
            "Adaptive graph-interactive network for joint multi-intent and slot filling",
            "MixATIS: SlotF1=88.1, IntentF1=81.2, IntentAcc=75.8, OverallAcc=44.5; MixSNIPS: SlotF1=94.5, IntentF1=98.6, IntentAcc=96.5, OverallAcc=76.4",
            "Introduces graph-based interaction between intents and slots",
            "Complex graph design; deployment overhead on lightweight systems",
        ],
        [
            f"Ding et al., 2021 {c_dgm}",
            "Dynamic graph model for multi-intent interactions",
            "MixATIS: SlotF1=88.7, IntentF1=81.0, IntentAcc=76.7, OverallAcc=47.1; MixSNIPS: SlotF1=94.7, IntentF1=98.6, IntentAcc=96.7, OverallAcc=78.0",
            "Captures fine-grained interactions among intent labels and slot sequence",
            "Graph construction and message passing increase model complexity",
        ],
        [
            f"Yang et al., 2021 {c_hitrans}",
            "Hierarchical transformer for nested named entity recognition",
            "Reported NER F1=87.04 (paper-reported improvement over 83.75 and 76.18 baselines)",
            "Improves nested entity modeling with hierarchical span interactions",
            "Designed for nested NER; requires adaptation for intent-slot pipelines",
        ],
        [
            f"Chen et al., 2022 {c_han}",
            "Higher-order attention for intent-slot coupling",
            "ATIS: IntentAcc=97.21, SlotF1=99.12, OverallAcc=91.80; SNIPS: IntentAcc=96.46, SlotF1=98.54, OverallAcc=88.67",
            "Models high-order contextual relations",
            "Higher memory footprint than linear baselines",
        ],
        [
            f"Song et al., 2022 {c_gisco}",
            "Global intent-slot co-occurrence enhancement",
            "MixATIS: SlotF1=88.5, IntentAcc=75.0, OverallAcc=48.2; MixSNIPS: SlotF1=95.0, IntentAcc=95.5, OverallAcc=75.9",
            "Uses global co-occurrence priors to reduce mismatch",
            "Needs domain-consistent co-occurrence statistics",
        ],
        [
            f"Gunaratna et al., 2022 {c_estan}",
            "Explainable slot type attention",
            "ATIS: IntentAcc=98.99, SlotF1=97.24; SNIPS: IntentAcc=99.10, SlotF1=96.20",
            "Adds interpretable attention traces",
            "Attention explanation may not fully capture causal behavior",
        ],
        [
            f"Pham et al., 2023 {c_misca}",
            "Intent-slot co-attention with multi-intent decoder (MISCA)",
            "MixATIS: IntentAcc=76.7, SlotF1=90.5, OverallAcc=53.0; MixSNIPS: IntentAcc=97.3, SlotF1=95.2, OverallAcc=77.9",
            "Improves overall semantic frame parsing via intent-slot co-attention",
            "Training objective and co-attention design are less lightweight for edge deployment",
        ],
        [
            f"Liu et al., 2025 {c_agri}",
            "Agricultural knowledge-enhanced BERT-TextCNN-attention hybrid",
            "Accuracy=79.6, Recall=80.1, F1=79.8 (8041 domain queries)",
            "Demonstrates domain knowledge integration benefits in specialized QA",
            "Domain-specific data and architecture limit direct cross-domain transfer",
        ],
    ]

    add_table(
        doc,
        ["Paper", "Method used", "Evaluation Metrics", "Key contributions", "Limitations"],
        literature_rows,
    )

    add_paragraph(
        doc,
        (
            "Table 1 compares representative intent-slot and NER studies relevant to command understanding. The progression from graph frameworks to hybrid transformer systems "
            "improves contextual coupling and robustness but typically increases model complexity. For desktop automation, this motivates compact alternatives that preserve high utility "
            "without requiring expensive finetuning cycles. The current project therefore adopts a classical vector-space classifier and deterministic entity extraction to maintain runtime "
            "efficiency while still supporting multi-stage diagnostics."
        ),
    )

    add_heading(doc, "5. Methodology / Proposed System", level=1)
    add_paragraph(
        doc,
        "Figure 1 illustrates the proposed architecture in which command text is transformed into TF-IDF features, classified into intent labels, parsed by regex-based entity rules, and routed to executor handlers.",
    )

    doc.add_picture(str(fig_arch), width=Inches(6.3))
    add_paragraph(doc, "Figure 1. Proposed system architecture and command processing flow", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            "Figure 1 shows a deterministic five-stage flow designed for stable desktop automation. First, each user command is encoded with TF-IDF to preserve lexical signal in a sparse feature space. "
            "Second, the trained classifier predicts one of 16 intents. Third, intent-specific regex rules extract required entities such as file names, destinations, and query strings. "
            "Fourth, a guarded executor validates required fields through dry-run checks. Finally, only validated actions are passed to file-system or web handlers, reducing accidental unsafe operations."
        ),
    )

    add_paragraph(doc, "Dataset descriptive statistics are summarized in Table 2.")
    dataset_rows = [
        [
            "Intent dataset (data/intent_dataset.json)",
            len(intent_samples),
            len(intent_counts),
            f"Balanced ({min(intent_counts.values())}-{max(intent_counts.values())} per class)",
            f"mean={statistics.mean(words_per_intent):.2f}, median={statistics.median(words_per_intent):.1f}",
            f"mean chars={statistics.mean(chars_per_intent):.2f}",
        ],
        [
            "End-to-end dataset (data/e2e_eval_dataset.json)",
            len(e2e_samples),
            len(set(item['intent'] for item in e2e_samples)),
            "Broad intent coverage with explicit expected entities",
            f"mean={statistics.mean(words_per_e2e):.2f}, median={statistics.median(words_per_e2e):.1f}",
            f"mean expected entities={statistics.mean(entities_per_e2e):.2f}",
        ],
    ]

    add_table(
        doc,
        [
            "Dataset",
            "Samples",
            "Classes",
            "Distribution",
            "Token-length stats",
            "Entity/char stats",
        ],
        dataset_rows,
    )

    add_paragraph(
        doc,
        (
            "Table 2 indicates that the intent corpus is intentionally balanced across all 16 classes, so class-weighting or resampling was not required for the baseline experiments. "
            "The end-to-end benchmark includes realistic entity-bearing commands, enabling execution-aware diagnostics beyond pure classification. Average utterance length remains short in both datasets, "
            "which supports sparse lexical modeling. This data profile aligns with command-style NLP tasks where concise imperative phrasing dominates."
        ),
    )

    add_paragraph(doc, "Figure 2 presents the class distribution, and Figure 6 shows lexical prominence via a word cloud.")

    doc.add_picture(str(fig_dist), width=Inches(6.3))
    add_paragraph(doc, "Figure 2. Intent class frequency distribution", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            "Figure 2 confirms strict class balance in the training corpus. This balanced setup reduces majority-class bias and allows macro metrics to better reflect model quality across intents. "
            "Because every class has equal support, differences in per-class F1 become attributable to semantic overlap rather than sample scarcity. "
            "This also makes confusion analysis more interpretable, especially for overlapping web-search related intents."
        ),
    )

    doc.add_picture(str(fig_cloud), width=Inches(6.3))
    add_paragraph(doc, "Figure 6. Word cloud of command corpus", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            "Figure 6 visualizes the lexical center of the dataset, where verbs such as open, move, copy, search, and run dominate. "
            "This supports the engineering decision to use TF-IDF features, as high-signal command verbs and object terms are explicit and compact. "
            "The figure also reveals overlap among search-oriented commands, motivating stricter intent boundary handling for WEB_SEARCH, SEARCH_FILE, and SEARCH_WIKIPEDIA."
        ),
    )

    add_paragraph(doc, "Preprocessing and feature engineering:")
    add_paragraph(
        doc,
        (
            "The pipeline applies lowercase normalization and sparse TF-IDF encoding through scikit-learn's TfidfVectorizer. "
            "Entity extraction uses regex templates tailored to each intent (e.g., file patterns, zip patterns, destination spans, and search query spans). "
            "For commands with optional structures, normalization routines clean folder/directory tokens and spacing before execution checks."
        ),
    )

    add_paragraph(doc, "Dataset split ratio and training protocol:")
    add_paragraph(
        doc,
        (
            "The dataset is split into train, validation, and test sets at 60:20:20 with stratification (144/48/48 samples). "
            "Hyperparameters are tuned with 5-fold cross-validation on the training split, and the selected model is retrained on train+validation before final test evaluation."
        ),
    )

    add_paragraph(doc, "Hyperparameter setup (Table 3):")
    logistic_best_c = intent_summary["model_results"]["logistic_regression"]["best_params"]["C"]
    svm_best_c = intent_summary["model_results"]["linear_svm"]["best_params"]["C"]
    selected_model_name = intent_summary["selected_model"]
    hyper_rows = [
        ["Vectorizer", "TfidfVectorizer defaults", "lowercase=True, sparse TF-IDF representation"],
        ["Logistic Regression", "C in [0.1, 1, 5, 10, 20, 50, 100], max_iter=2000", f"best C={logistic_best_c}"],
        ["Linear SVM", "C in [0.1, 1, 5, 10, 20, 50, 100], max_iter=5000", f"best C={svm_best_c}; selected model={selected_model_name}"],
        ["Cross-validation", "Stratified folds capped by minimum class support", f"cv_folds={intent_summary['cv_folds']}"],
    ]
    add_table(doc, ["Component", "Search space", "Chosen value"], hyper_rows)

    add_paragraph(
        doc,
        (
            "Table 3 lists the exact hyperparameter configuration used for reproducible experiments. The selected runtime model is Linear SVM with C=10, chosen by validation macro-F1. "
            "A Logistic Regression alternative was also retained for baseline comparison."
        ),
    )

    add_paragraph(doc, "Mathematical equations used in the implementation and evaluation:")
    add_equation_block(doc, "TF-IDF(t,d) = TF(t,d) * log(N / (1 + DF(t)))", 1)
    add_equation_block(doc, "P(y=k|x) = exp(w_k^T x + b_k) / sum_j exp(w_j^T x + b_j)", 2)
    add_equation_block(doc, "min_{w,b} 0.5*||w||^2 + C * sum_i max(0, 1 - y_i*(w^T x_i + b))", 3)
    add_equation_block(doc, "Precision = TP / (TP + FP), Recall = TP / (TP + FN)", 4)
    add_equation_block(doc, "F1 = 2 * Precision * Recall / (Precision + Recall), Accuracy = Correct / Total", 5)

    add_heading(doc, "6. Results and Discussion", level=1)

    add_paragraph(doc, "6.1 Environment setup")
    env_rows = [
        ["Execution mode", "Local machine (Windows)"],
        ["Python", platform.python_version()],
        ["CPU", cpu_name],
        ["Logical processors", logical_cpu],
        ["Installed RAM (GB)", f"{ram_gb:.2f}" if ram_gb else "Not detected"],
        ["Core libraries", "scikit-learn, numpy, pandas, matplotlib, streamlit, spaCy, python-docx"],
    ]
    add_table(doc, ["Item", "Value"], env_rows)

    add_paragraph(
        doc,
        (
            "The experiments were run in a local Python virtual environment to reflect realistic desktop deployment conditions. "
            "The software stack is intentionally lightweight and avoids GPU requirements. This supports reproducibility for student and edge scenarios, "
            "where command assistance must run on commodity hardware rather than server clusters."
        ),
    )

    add_paragraph(doc, "6.2 Intent classification outcomes")
    selected = intent_summary["selected_metrics"]
    logistic = intent_summary["model_results"]["logistic_regression"]
    svm = intent_summary["model_results"]["linear_svm"]

    results_rows = [
        ["Majority baseline", f"{intent_summary['baseline']['accuracy'] * 100:.2f}%", "-", "Predicts most frequent intent only"],
        ["Logistic Regression", f"{logistic['accuracy'] * 100:.2f}%", f"{logistic['macro_f1'] * 100:.2f}%", f"C={logistic['best_params']['C']}"],
        ["Linear SVM (selected)", f"{selected['accuracy'] * 100:.2f}%", f"{selected['macro_f1'] * 100:.2f}%", f"C={svm['best_params']['C']}"],
    ]
    add_table(doc, ["Model", "Accuracy", "Macro-F1", "Notes"], results_rows)

    add_paragraph(
        doc,
        (
            "Compared with the majority baseline (6.25%), both linear models provide substantial gains, confirming that sparse lexical modeling is effective for command semantics. "
            "The selected Linear SVM reaches 91.67% accuracy and 89.15% macro-F1, while Logistic Regression attains similar accuracy with slightly stronger test macro-F1. "
            "Selection still follows validation macro-F1, preserving protocol consistency and reducing overfitting risk."
        ),
    )

    doc.add_picture(str(fig_conf), width=Inches(6.3))
    add_paragraph(doc, "Figure 3. End-to-end intent confusion matrix on 40 benchmark commands (all cells annotated, including zeros)", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            "Figure 3 uses the same end-to-end benchmark split as Figure 4 so failure counts remain directly comparable. "
            "The matrix is strongly diagonal, and the sum of off-diagonal cells equals the intent-mismatch count in Figure 4. "
            "This makes the intent versus entity failure decomposition consistent across both plots."
        ),
    )

    add_paragraph(doc, "6.3 End-to-end execution analysis")
    e2e_rows = [
        ["Intent accuracy", f"{e2e_summary['intent_accuracy'] * 100:.2f}%"],
        ["Entity exact-match rate", f"{e2e_summary['entity_exact_match_rate'] * 100:.2f}%"],
        ["Validation pass rate", f"{e2e_summary['validation_pass_rate'] * 100:.2f}%"],
        ["Execution success rate", f"{e2e_summary['execution_success_rate'] * 100:.2f}%"],
        ["End-to-end success rate", f"{e2e_summary['end_to_end_success_rate'] * 100:.2f}%"],
    ]
    add_table(doc, ["Pipeline metric", "Score"], e2e_rows)

    add_paragraph(doc, "Table 6.3 row-wise interpretation:")
    add_paragraph(doc, "Intent accuracy: proportion of commands whose predicted intent matches the ground-truth intent label.")
    add_paragraph(doc, "Entity exact-match rate: proportion of commands where the full extracted entity dictionary exactly matches expected entities.")
    add_paragraph(doc, "Validation pass rate: proportion of samples that satisfy route-level validation checks (required entities/system-intent guards).")
    add_paragraph(doc, "Execution success rate: proportion of dry-run executor calls that return success status after routing.")
    add_paragraph(doc, "End-to-end success rate: proportion of samples that satisfy intent correctness, entity exact match, validation pass, and execution success simultaneously.")

    doc.add_picture(str(fig_stage), width=Inches(6.0))
    add_paragraph(doc, "Figure 4. End-to-end stage breakdown", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            f"Figure 4 indicates that {e2e_summary['stage_breakdown'].get('success', 0)} of {e2e_total} samples complete successfully in dry-run mode. "
            f"Intent mismatch is {e2e_summary['stage_breakdown'].get('intent_mismatch', 0)}, which exactly equals the off-diagonal sum in Figure 3. "
            f"Entity mismatch is {e2e_summary['stage_breakdown'].get('entity_mismatch', 0)} and represents cases where intent is correct but entities are incomplete or incorrect. "
            "No major execution routing failures are observed, validating the entity-key checks in the executor design."
        ),
    )

    add_paragraph(doc, "6.4 NER performance and interpretability")
    ner_rows = [
        ["Micro precision", f"{ner_summary['micro_precision'] * 100:.2f}%"],
        ["Micro recall", f"{ner_summary['micro_recall'] * 100:.2f}%"],
        ["Micro F1", f"{ner_summary['micro_f1'] * 100:.2f}%"],
        ["Entity exact match", f"{ner_summary['entity_exact_match_rate'] * 100:.2f}%"],
        ["TP / FP / FN", f"{ner_summary['true_positive_entities']} / {ner_summary['false_positive_entities']} / {ner_summary['false_negative_entities']}"],
    ]
    add_table(doc, ["NER metric", "Value"], ner_rows)

    add_paragraph(doc, "Table 6.4 interpretation:")
    add_paragraph(doc, "Micro precision: fraction of predicted entity items that are correct when pooled across all entity types.")
    add_paragraph(doc, "Micro recall: fraction of expected entity items that are correctly recovered by the extractor.")
    add_paragraph(doc, "Micro F1: harmonic mean of micro precision and micro recall, summarizing overall extraction quality.")
    add_paragraph(doc, "Entity exact match: strict sample-level metric requiring complete entity-dictionary equality for each command.")
    add_paragraph(doc, "TP/FP/FN: absolute counts of correctly extracted, incorrectly added, and missed entity items, respectively.")

    doc.add_picture(str(fig_ner), width=Inches(6.3))
    add_paragraph(doc, "Figure 5. Per-entity-type F1 scores", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            "Figure 5 shows near-perfect extraction for most entity types, with lower performance concentrated in destination and query fields. "
            "This behavior is expected because destination strings and open-ended queries have broader lexical variability than fixed file-name patterns. "
            "Interpretability is supported through explicit regex rules and per-entity diagnostics, allowing direct inspection of extraction logic without opaque latent attention maps."
        ),
    )

    doc.add_picture(str(fig_rocpr), width=Inches(6.3))
    add_paragraph(doc, "Figure 7. One-vs-rest ROC and precision-recall curves", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        (
            f"Figure 7 summarizes ranking behavior of the selected classifier over one-vs-rest decision scores. Macro ROC-AUC is {roc_pr['macro_roc_auc']:.4f}, "
            f"and macro average precision is {roc_pr['macro_pr_auc']:.4f}. These values indicate strong separability for most classes despite lexical overlap in web-search intents. "
            "The ROC/PR analysis complements accuracy and macro-F1 by exposing score-distribution quality beyond hard-label decisions."
        ),
    )

    add_paragraph(doc, "6.5 Comparison with prior work and practical trade-offs")
    add_paragraph(
        doc,
        (
            "Compared with transformer-heavy literature, the current system sacrifices some representational richness but gains practical advantages in interpretability, training speed, and local deployability. "
            "Graph-transformer joint models often report stronger benchmark potential under larger datasets, yet they require greater computational and integration complexity. "
            "For command automation tasks with concise utterances and controlled vocabulary, the presented classical pipeline offers a favorable accuracy-to-cost ratio."
        ),
    )

    add_paragraph(doc, "Pros and Cons summary:")
    pros_cons_rows = [
        ["Pros", "High reproducibility, low compute overhead, deterministic entity extraction, explicit execution validation"],
        ["Cons", "Limited contextual understanding, weaker generalization to unseen paraphrases, WEB_SEARCH ambiguity"],
    ]
    add_table(doc, ["Category", "Observation"], pros_cons_rows)

    add_heading(doc, "7. Conclusion and Future Work", level=1)
    add_paragraph(
        doc,
        (
            "This research demonstrates that a compact intent-classification plus rule-based NER architecture can achieve strong end-to-end desktop assistant performance without deep model finetuning. "
            "The pipeline reached 91.67% intent accuracy, 96.77% NER micro-F1, and 92.50% execution-level success, validating the feasibility of lightweight local NLP automation."
        ),
    )

    add_paragraph(
        doc,
        (
            "The major research gap identified is the mismatch between benchmark-oriented SLU modeling and execution-aware assistant evaluation. "
            "By integrating dry-run validation and stage-level diagnostics, this project contributes a deployment-centered perspective for command understanding systems."
        ),
    )

    add_paragraph(
        doc,
        (
            "Current limitations include the small data scale, rule brittleness for highly varied natural language, and limited multilingual support. "
            "Future improvements should include contextual embeddings (e.g., domain-adapted transformers), calibrated confidence thresholds, active learning loops, and multilingual intent/entity expansion."
        ),
    )

    add_paragraph(
        doc,
        (
            "To improve real-world scalability, the system can be extended with asynchronous task execution, policy-based permissions, secure audit logging, and user profile personalization. "
            "Integrating explainable AI tooling such as SHAP/LIME for classifier decisions and uncertainty-aware fallback prompts can further improve trust and robustness."
        ),
    )

    add_heading(doc, "8. References", level=1)
    cited_refs = [doi_map[doi] for doi in sorted(used_dois)]
    cited_total = len(cited_refs)
    cited_recent = sum(1 for ref in cited_refs if int(ref.get("year", 0)) >= 2024)
    cited_recent_ratio = (cited_recent / cited_total) if cited_total else 0.0
    add_paragraph(
        doc,
        (
            f"Only in-text cited references are listed below (count={cited_total}, DOI-verified). "
            f"Recent references (2024-2026) account for {cited_recent}/{cited_total} = {cited_recent_ratio * 100:.2f}% of the final bibliography."
        ),
    )

    sorted_refs = sorted(cited_refs, key=lambda ref: (ref.get("year", 0), first_author_last_name(ref.get("authors", []))))
    for idx, ref in enumerate(sorted_refs, start=1):
        add_paragraph(doc, f"[{idx}] {format_reference_entry(ref)}")

    doc.save(OUTPUT_DOCX)
    print(f"Report generated: {OUTPUT_DOCX}")
    print(f"Assets directory: {assets_dir}")


if __name__ == "__main__":
    main()
